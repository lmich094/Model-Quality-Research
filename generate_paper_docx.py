#!/usr/bin/env python3
"""Generate results/research_paper.docx from results/research_paper.md using python-docx."""

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


REPO_ROOT = Path(__file__).parent
MD_PATH = REPO_ROOT / "results" / "research_paper.md"
DOCX_PATH = REPO_ROOT / "results" / "research_paper.docx"


def add_formatted_runs(paragraph, text: str):
    """Parse **bold** and *italic* markers and add styled runs to a paragraph."""
    # Replace escaped pipes so they survive the run splitting
    text = text.replace("\\|", "\x00PIPE\x00")

    pattern = r"(\*\*[^*]+\*\*|\*[^*]+\*|[^*]+|\*)"
    for match in re.finditer(pattern, text):
        segment = match.group(0)
        segment = segment.replace("\x00PIPE\x00", "|")
        if segment.startswith("**") and segment.endswith("**") and len(segment) > 4:
            run = paragraph.add_run(segment[2:-2])
            run.bold = True
        elif segment.startswith("*") and segment.endswith("*") and len(segment) > 2:
            run = paragraph.add_run(segment[1:-1])
            run.italic = True
        else:
            paragraph.add_run(segment)


def is_separator_row(row: str) -> bool:
    """Return True if this is a markdown table separator row (|---|---|)."""
    return bool(re.match(r"^\|[-| :]+\|$", row.strip()))


def split_table_row(row: str):
    """Split a markdown table row into cells, handling escaped pipes."""
    safe = row.strip().replace("\\|", "\x00PIPE\x00")
    cells = [c.replace("\x00PIPE\x00", "|").strip() for c in safe.split("|")[1:-1]]
    return cells


def collect_table(lines, start: int):
    """Collect consecutive table lines starting at start. Returns (rows, next_index)."""
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        line = lines[i].strip()
        if is_separator_row(line):
            i += 1
            continue
        rows.append(split_table_row(line))
        i += 1
    return rows, i


def add_table(doc: Document, rows: list):
    if not rows:
        return
    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = "Table Grid"

    for row_idx, row_data in enumerate(rows):
        tr = table.rows[row_idx]
        for col_idx in range(num_cols):
            cell_text = row_data[col_idx] if col_idx < len(row_data) else ""
            cell = tr.cells[col_idx]
            cell.text = ""
            para = cell.paragraphs[0]
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after = Pt(2)
            if row_idx == 0:
                run = para.add_run(cell_text)
                run.bold = True
            else:
                add_formatted_runs(para, cell_text)


def set_style_properties(doc: Document):
    """Adjust built-in styles for a clean, readable look."""
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    for level, size, bold in [
        ("Heading 1", 16, True),
        ("Heading 2", 13, True),
        ("Heading 3", 12, True),
    ]:
        style = doc.styles[level]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor(0x1F, 0x39, 0x64)


def convert(md_path: Path, docx_path: Path):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    set_style_properties(doc)

    lines = md_path.read_text(encoding="utf-8").splitlines()

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Blank line
        if not stripped:
            i += 1
            continue

        # Horizontal rule
        if stripped in ("---", "***", "___"):
            p = doc.add_paragraph()
            run = p.add_run("─" * 60)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            i += 1
            continue

        # H1 — document title
        if re.match(r"^# (?!#)", stripped):
            text = stripped[2:].strip()
            p = doc.add_heading(text, level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # H2 — subtitle (treated as H1 visually if it starts with "##" not "###")
        if re.match(r"^## (?!#)", stripped):
            text = stripped[3:].strip()
            p = doc.add_heading(text, level=1)
            i += 1
            continue

        # H3
        if re.match(r"^### (?!#)", stripped):
            text = stripped[4:].strip()
            doc.add_heading(text, level=2)
            i += 1
            continue

        # H4
        if re.match(r"^#### (?!#)", stripped):
            text = stripped[5:].strip()
            doc.add_heading(text, level=3)
            i += 1
            continue

        # Table
        if stripped.startswith("|"):
            rows, i = collect_table(lines, i)
            add_table(doc, rows)
            doc.add_paragraph()  # spacing after table
            continue

        # Unordered list item
        if re.match(r"^[-*] ", stripped):
            text = stripped[2:].strip()
            p = doc.add_paragraph(style="List Bullet")
            add_formatted_runs(p, text)
            i += 1
            continue

        # Numbered list item
        if re.match(r"^\d+\. ", stripped):
            text = re.sub(r"^\d+\. ", "", stripped)
            p = doc.add_paragraph(style="List Number")
            add_formatted_runs(p, text)
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        add_formatted_runs(p, stripped)
        i += 1

    doc.save(docx_path)
    print(f"Saved: {docx_path}")


if __name__ == "__main__":
    if not MD_PATH.exists():
        print(f"Error: {MD_PATH} not found", file=sys.stderr)
        sys.exit(1)
    convert(MD_PATH, DOCX_PATH)
