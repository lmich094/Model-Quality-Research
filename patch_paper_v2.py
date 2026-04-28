#!/usr/bin/env python3
"""
Patch "Regenerative Ag Research Paper.docx" with the agreed edits and save
as "Regenerative Ag Research Paper v2.docx".

Changes applied:
  1. Abstract: add κ pointer to Section 3.6 at first mention
  2. Section 2.4: clarify n=208 paired rows (not 418)
  3. Section 3.1: define "pooled models" before Table 1
  4. Section 3.5: spell out "moderate framing level (MOD)" instead of "MOD framing gap"
  5. Section 3.6: insert Cohen's Kappa definition before Table 5
  6. Section 4.1 H5: expand explanation of what "framing" means and where H5 holds
"""

import copy
import shutil
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

REPO = Path(__file__).parent
SRC = REPO / "results" / "Regenerative Ag Research Paper.docx"
DST = REPO / "results" / "Regenerative Ag Research Paper v2.docx"

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def replace_text_in_para(para, old: str, new: str):
    """Replace old with new inside a paragraph, preserving other run text."""
    full = para.text
    if old not in full:
        return False
    new_full = full.replace(old, new, 1)
    # Wipe all runs and write the new full text into the first run
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = new_full
    else:
        para.add_run(new_full)
    return True


def insert_paragraph_before(ref_para, text: str):
    """Insert a Normal-styled paragraph containing text immediately before ref_para."""
    new_p = OxmlElement("w:p")

    # Copy paragraph properties (style) from the reference paragraph
    ref_pPr = ref_para._element.find(qn("w:pPr"))
    if ref_pPr is not None:
        new_p.append(copy.deepcopy(ref_pPr))

    new_r = OxmlElement("w:r")
    new_t = OxmlElement("w:t")
    new_t.text = text
    new_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    new_r.append(new_t)
    new_p.append(new_r)

    ref_para._element.addprevious(new_p)


def insert_paragraph_after(ref_para, text: str):
    """Insert a Normal-styled paragraph containing text immediately after ref_para."""
    new_p = OxmlElement("w:p")

    ref_pPr = ref_para._element.find(qn("w:pPr"))
    if ref_pPr is not None:
        new_p.append(copy.deepcopy(ref_pPr))

    new_r = OxmlElement("w:r")
    new_t = OxmlElement("w:t")
    new_t.text = text
    new_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    new_r.append(new_t)
    new_p.append(new_r)

    ref_para._element.addnext(new_p)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

shutil.copy(SRC, DST)
doc = Document(DST)

# We'll do in-place text replacements first, then insertions.
# Re-fetch paragraphs each time we insert to avoid stale references.

# ── Change 1: Abstract – add κ pointer ──────────────────────────────────────
OLD1 = "the dimension most central to the study's question."
NEW1 = "the dimension most central to the study's question. (κ is Cohen's Kappa; see Section 3.6 for a full explanation.)"

# ── Change 2: Section 2.4 – clarify n=208 ───────────────────────────────────
OLD2 = "N = 208 paired rows (two response sets × 104 rows each, after excluding the Q7-PRO-S error row from both sets)."
NEW2 = (
    "N = 208 paired rows — meaning 208 individual AI responses, each scored "
    "independently by both raters. Rater scores appear as two columns on the same "
    "row, not as separate rows, so the IRA dataset is 208 rows, not 418. For "
    "reference: the raw response count is 104 Claude + 105 Codex = 209 responses; "
    "Q7-PRO-S is excluded from both rater sets, leaving 208."
)

# ── Change 4: Section 3.5 – spell out moderate framing level ────────────────
OLD4 = "the MOD framing gap closes to zero entirely"
NEW4 = "the gap between Claude and Codex at the moderate framing level (MOD) closes to zero entirely"

# ── Change 6a: Section 4.1 H5 – replace existing paragraph text ─────────────
OLD6 = (
    "For RS across all conditions and for CC under most conditions, framing produces "
    "negligible differences. For composting specifically, framing is essentially "
    "irrelevant. The meaningful framing effects are real, but they're concentrated in "
    "specific dimensions (TD, APK) and vary substantially by domain."
)
NEW6_PARA1 = (
    'Here, "framing" refers specifically to the expertise prefix manipulation — the '
    "sentence added to the front of each prompt that declares who you are (novice, "
    "moderate, expert, etc.). H5 is the skeptical baseline: that those prefix sentences "
    "are essentially ignored, and the AI produces the same response regardless of "
    "whether you claim to be a complete beginner or a professional agronomist."
)
NEW6_PARA2 = (
    "The data partially bears this out. For RS, framing made almost no difference "
    "across the board — scores stayed within a 0.04-point range (2.81 to 2.85) no "
    "matter who was asking. For CC, novice framing produced essentially the same "
    "low-caveat responses as expert framing in most conditions. For the composting "
    "domain specifically, TD barely shifted at all across framing levels (range = 0.33 "
    "on a 5-point scale), meaning the AI gave essentially the same depth of response to "
    "a novice and a professional. So H5 holds in patches — the null result is real and "
    "worth acknowledging, even if TD and APK show that framing isn't completely "
    "ignored either."
)

# Apply text replacements
h5_para_ref = None
for para in doc.paragraphs:
    replace_text_in_para(para, OLD1, NEW1)
    replace_text_in_para(para, OLD2, NEW2)
    replace_text_in_para(para, OLD4, NEW4)
    if OLD6 in para.text:
        replace_text_in_para(para, OLD6, NEW6_PARA1)
        h5_para_ref = para

# Insert the second H5 paragraph after the first
if h5_para_ref is not None:
    insert_paragraph_after(h5_para_ref, NEW6_PARA2)

# ── Change 3: Section 3.1 – insert pooled models definition before Table 1 ──
POOLED_DEF = (
    'Throughout Sections 3.1, 3.2, and 3.4, "pooled models" means Claude and Codex '
    "responses are combined into a single dataset — 104 Claude rows and 105 Codex rows "
    "treated as one group — to isolate the effect of the variable being examined "
    "(framing or length) from model-level differences. Model-specific differences are "
    "reported separately in Section 3.3."
)
TABLE1_CAPTION = "Table 1. Mean scores by expertise framing"

for para in doc.paragraphs:
    if para.text.strip().startswith(TABLE1_CAPTION):
        insert_paragraph_before(para, POOLED_DEF)
        break

# ── Change 5: Section 3.6 – insert Cohen's Kappa definition before Table 5 ──
KAPPA_DEF = (
    "Cohen's Kappa (κ) is a standard measure of agreement between two raters that "
    "corrects for chance. If two raters would agree 60% of the time just by guessing "
    "randomly, a raw 70% agreement rate isn't impressive — κ accounts for that baseline "
    "and reports only the agreement above chance. κ = 1.0 means perfect agreement; "
    "κ = 0.0 means agreement is no better than chance; negative κ means raters disagree "
    "more than chance would predict. Conventional benchmarks: < 0.20 slight, 0.21–0.40 "
    "fair, 0.41–0.60 moderate, 0.61–0.80 substantial, > 0.80 near-perfect. "
    "Linear-weighted κ (used here) gives partial credit for near-misses, so a 1-point "
    "disagreement is penalized less than a 3-point disagreement."
)
TABLE5_CAPTION = "Table 5. Inter-rater reliability"

for para in doc.paragraphs:
    if para.text.strip().startswith(TABLE5_CAPTION):
        insert_paragraph_before(para, KAPPA_DEF)
        break

doc.save(DST)
print(f"Saved: {DST}")
